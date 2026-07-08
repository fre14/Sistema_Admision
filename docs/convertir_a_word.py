#!/usr/bin/env python3
"""
Convierte los archivos Markdown del trabajo monográfico SAA a un documento Word (.docx)
con formato académico profesional.
"""
import re
import os
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
OUTPUT = os.path.join(DOCS_DIR, "Trabajo_Monografico_SAA.docx")

FILES = [
    os.path.join(DOCS_DIR, "capitulos_I_II_III.md"),
    os.path.join(DOCS_DIR, "capitulo_IV.md"),
    os.path.join(DOCS_DIR, "capitulo_V_refs_anexos.md"),
]

def setup_styles(doc):
    """Configura los estilos del documento."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.space_after = Pt(6)
    
    # Configurar margenes
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.54)

def add_cover_page(doc):
    """Agrega la portada del documento."""
    # Espaciado superior
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
    
    # Año
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AÑO DE LA ESPERANZA Y EL FORTALECIMIENTO DE LA DEMOCRACIA")
    run.bold = True
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Universidad
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("UNIVERSIDAD NACIONAL DE SAN CRISTÓBAL DE HUAMANGA")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("FACULTAD DE INGENIERÍA DE MINAS, GEOLOGÍA Y CIVIL")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ESCUELA PROFESIONAL DE INGENIERÍA DE SISTEMAS")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    for _ in range(3):
        doc.add_paragraph()
    
    # Tipo de trabajo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("TRABAJO MONOGRÁFICO")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = 'Times New Roman'
    
    doc.add_paragraph()
    
    # Título
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("SISTEMA AUTOMATIZADO DE ADMISIÓN (SAA) BASADO EN EL ENFOQUE DE DESARROLLO DIRIGIDO POR ESPECIFICACIONES (SSD) PARA LA GESTIÓN DEL PROCESO DE EVALUACIÓN Y SELECCIÓN DE CANDIDATOS")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    
    for _ in range(3):
        doc.add_paragraph()
    
    # Curso
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CURSO: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run("Pruebas y Aseguramiento de la Calidad (IS-489)")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # Docente
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("DOCENTE: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run("Ing. Richard Zapata Casaverde")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    # Alumno
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("ALUMNO: ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    run = p.add_run("Fredy Bonilla Rey")
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    for _ in range(4):
        doc.add_paragraph()
    
    # Lugar y año
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AYACUCHO – PERÚ")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'
    
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("2026")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = 'Times New Roman'
    
    doc.add_page_break()

def clean_markdown(text):
    """Limpia formato markdown de una línea de texto."""
    # Remove bold markers
    text = re.sub(r'\*\*\*(.*?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    # Remove inline code
    text = re.sub(r'`(.*?)`', r'\1', text)
    # Remove links [text](url)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    # Remove HTML tags
    text = re.sub(r'<[^>]+>', '', text)
    # Remove > blockquotes
    text = re.sub(r'^>\s*', '', text)
    # Remove [!NOTE], [!IMPORTANT], etc.
    text = re.sub(r'\[!(NOTE|TIP|IMPORTANT|WARNING|CAUTION)\]', '', text)
    return text.strip()

def add_formatted_run(paragraph, text, bold=False, italic=False, size=12):
    """Agrega un run formateado a un párrafo."""
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run

def parse_table(lines, start_idx):
    """Parsea una tabla markdown y retorna las filas y el índice final."""
    rows = []
    i = start_idx
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith('|'):
            break
        # Skip separator lines
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        cells = [clean_markdown(c.strip()) for c in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
        i += 1
    return rows, i

def add_table_to_doc(doc, rows):
    """Agrega una tabla al documento Word."""
    if not rows:
        return
    
    num_cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols, style='Table Grid')
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    for i, row in enumerate(rows):
        for j, cell_text in enumerate(row):
            if j < num_cols:
                cell = table.cell(i, j)
                cell.text = ''
                p = cell.paragraphs[0]
                run = p.add_run(cell_text)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                if i == 0:  # Header row
                    run.bold = True
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
    
    doc.add_paragraph()  # Space after table

def process_markdown_file(doc, filepath):
    """Procesa un archivo markdown y lo agrega al documento Word."""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()
    
    lines = content.split('\n')
    i = 0
    in_code_block = False
    code_lines = []
    skip_cover = True  # Skip the cover page section in first file
    
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()
        
        # Skip empty divs and HTML
        if stripped.startswith('<div') or stripped.startswith('</div') or stripped == '---':
            i += 1
            continue
        
        # Skip page break markers
        if 'page-break' in stripped:
            i += 1
            continue
        
        # Handle code blocks
        if stripped.startswith('```'):
            if in_code_block:
                # End code block - add accumulated code
                code_text = '\n'.join(code_lines)
                if code_text.strip():
                    p = doc.add_paragraph()
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run(code_text)
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
                code_lines = []
            i += 1
            continue
        
        if in_code_block:
            code_lines.append(line)
            i += 1
            continue
        
        # Handle tables
        if stripped.startswith('|') and i + 1 < len(lines) and lines[i + 1].strip().startswith('|'):
            rows, new_i = parse_table(lines, i)
            if rows:
                add_table_to_doc(doc, rows)
            i = new_i
            continue
        
        # Handle headers
        if stripped.startswith('#'):
            header_match = re.match(r'^(#{1,5})\s+(.+)', stripped)
            if header_match:
                level = len(header_match.group(1))
                header_text = clean_markdown(header_match.group(2))
                
                # Skip the cover page headers in first file
                if skip_cover and ('AÑO DE LA' in header_text or 'UNIVERSIDAD' in header_text or 
                    'FACULTAD' in header_text or 'ESCUELA' in header_text or
                    'TRABAJO MONOGRÁFICO' in header_text or 'SISTEMA AUTOMATIZADO' in header_text):
                    i += 1
                    continue
                
                if 'INTRODUCCIÓN' in header_text or 'CAPÍTULO' in header_text:
                    skip_cover = False
                
                # Map markdown header levels to Word heading styles
                if level == 1:
                    p = doc.add_heading(header_text, level=1)
                elif level == 2:
                    p = doc.add_heading(header_text, level=2)
                elif level == 3:
                    p = doc.add_heading(header_text, level=3)
                elif level == 4:
                    p = doc.add_heading(header_text, level=4)
                else:
                    p = doc.add_paragraph()
                    run = p.add_run(header_text)
                    run.bold = True
                    run.font.size = Pt(11)
                    run.font.name = 'Times New Roman'
                
                i += 1
                continue
        
        # Handle bullet points
        bullet_match = re.match(r'^(\s*)[-*]\s+(.+)', stripped)
        if bullet_match:
            text = clean_markdown(bullet_match.group(2))
            if text:
                p = doc.add_paragraph(text, style='List Bullet')
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            i += 1
            continue
        
        # Handle numbered lists  
        numbered_match = re.match(r'^(\s*)\d+[\.\)]\s+(.+)', stripped)
        if numbered_match:
            text = clean_markdown(numbered_match.group(2))
            if text:
                p = doc.add_paragraph(text, style='List Number')
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(12)
            i += 1
            continue
        
        # Handle regular paragraphs
        text = clean_markdown(stripped)
        if text and len(text) > 3:  # Skip very short lines
            # Check if it's a label like "CURSO:", "DOCENTE:", "ALUMNO:"
            if any(label in text for label in ['CURSO:', 'DOCENTE:', 'ALUMNO:', 'AYACUCHO', '2026']):
                if skip_cover:
                    i += 1
                    continue
            
            p = doc.add_paragraph()
            p.paragraph_format.first_line_indent = Cm(1.25)
            
            # Handle bold segments within text
            parts = re.split(r'(\*\*.*?\*\*)', line.strip())
            for part in parts:
                part_clean = clean_markdown(part)
                if not part_clean:
                    continue
                is_bold = part.strip().startswith('**') and part.strip().endswith('**')
                run = p.add_run(part_clean)
                run.font.name = 'Times New Roman'
                run.font.size = Pt(12)
                run.bold = is_bold
        
        i += 1

def main():
    print("=" * 60)
    print("Generando documento Word del Trabajo Monográfico SAA...")
    print("=" * 60)
    
    doc = Document()
    setup_styles(doc)
    
    # 1. Portada
    print("[1/4] Generando portada...")
    add_cover_page(doc)
    
    # 2. Procesar cada archivo
    for idx, filepath in enumerate(FILES):
        filename = os.path.basename(filepath)
        print(f"[{idx+2}/4] Procesando {filename}...")
        if os.path.exists(filepath):
            process_markdown_file(doc, filepath)
            if idx < len(FILES) - 1:
                doc.add_page_break()
        else:
            print(f"  ⚠️ Archivo no encontrado: {filepath}")
    
    # 3. Guardar
    doc.save(OUTPUT)
    print(f"\n{'=' * 60}")
    print(f"✅ Documento generado exitosamente:")
    print(f"   {OUTPUT}")
    print(f"{'=' * 60}")

if __name__ == '__main__':
    main()

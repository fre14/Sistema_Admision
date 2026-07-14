#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Genera la seccion 4.1.5 Despliegue en Capa Gratuita del Trabajo Monografico SAA.
Incluye capturas reales del proceso de despliegue con sus explicaciones academicas.
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUTPUT = r"C:\Users\usuario\Downloads\Seccion_Despliegue_SAA.docx"

IMGS = {
    "render_crear_servicio": r"C:\Users\usuario\Downloads\IMAGENES DESPLIEGUE\WhatsApp Image 2026-07-08 at 5.29.03 PM.....jpeg",
    "render_static_vs_web":  r"C:\Users\usuario\Downloads\IMAGENES DESPLIEGUE\WhatsApp Image 2026-07-08 at 5.37.07 PM........jpeg",
    "render_configurar":     r"C:\Users\usuario\Downloads\IMAGENES DESPLIEGUE\WhatsApp Image 2026-07-08 at 5.37.09 PM.jpeg",
    "render_building":       r"C:\Users\usuario\Downloads\IMAGENES DESPLIEGUE\WhatsApp Image 2026-07-08 at 5.37.54 PM...jpeg",
    "render_live_logs":      r"C:\Users\usuario\Downloads\IMAGENES DESPLIEGUE\WhatsApp Image 2026-07-08 at 5.46.59 PM....jpeg",
    "somee_bd":              r"C:\Users\usuario\Downloads\IMAGENES DESPLIEGUE\WhatsApp Image 2026-07-09 at 12.17.28 AM.......jpeg",
    "vercel_import":         r"C:\Users\usuario\Downloads\IMAGENES DESPLIEGUE\WhatsApp Image 2026-07-08 at 5.19.24 PM.jpeg",
    "vercel_root_dir":       r"C:\Users\usuario\Downloads\IMAGENES DESPLIEGUE\WhatsApp Image 2026-07-08 at 5.20.41 PM..jpeg",
    "vercel_success":        r"C:\Users\usuario\Downloads\IMAGENES DESPLIEGUE\WhatsApp Image 2026-07-08 at 5.22.42 PM...jpeg",
}

fig_counter = [11]  # Continua la numeracion del documento principal (ajustar si es necesario)

def next_fig():
    n = fig_counter[0]
    fig_counter[0] += 1
    return n

def set_font(run, name='Times New Roman', size=12, bold=False, italic=False, color=None):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    sizes = {1: 14, 2: 13, 3: 12}
    set_font(run, size=sizes.get(level, 12), bold=True)
    return p

def body(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = Pt(21)  # 1.5
    run = p.add_run(text)
    set_font(run)
    return p

def insert_image(doc, img_path, caption_text, fig_num, width_cm=14):
    if not os.path.exists(img_path):
        body(doc, f"[IMAGEN NO ENCONTRADA: {img_path}]")
        return
    # Figura label encima
    p_label = doc.add_paragraph()
    p_label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_label.paragraph_format.space_before = Pt(6)
    r = p_label.add_run(f"Figura {fig_num}")
    set_font(r, bold=True, size=11)

    # Imagen
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(img_path, width=Cm(width_cm))

    # Caption debajo en italica
    p_cap = doc.add_paragraph()
    p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cap.paragraph_format.space_after = Pt(4)
    r_cap = p_cap.add_run(caption_text)
    set_font(r_cap, italic=True, size=10)

    # Nota
    p_nota = doc.add_paragraph()
    p_nota.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_nota.paragraph_format.space_after = Pt(12)
    r_nota = p_nota.add_run("Nota: Elaboracion propia.")
    set_font(r_nota, size=10)

def build_document():
    doc = Document()

    # Configurar margenes
    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.54)

    # Estilo normal
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # =========================================================
    # TITULO SECCION
    # =========================================================
    heading(doc, "4.1.5. Resultados de la Fase de Despliegue en Capa Gratuita", level=2)

    body(doc,
        "Una vez concluida la fase de implementacion y validacion del Sistema Automatizado de Admision (SAA), "
        "se procedio al despliegue del sistema en servicios de hospedaje gratuitos en la nube, como parte integral "
        "del ciclo de vida del software. Esta fase comprende tres componentes principales: (a) el despliegue del "
        "backend (API REST) desarrollado en ASP.NET Core 10 mediante la plataforma Render.com utilizando contenedores "
        "Docker; (b) el despliegue de la base de datos relacional SQL Server 2022 Express en la plataforma Somee.com; "
        "y (c) el despliegue del frontend desarrollado en React 19 mediante la plataforma Vercel. El cumplimiento de "
        "esta fase demuestra que el SAA es un sistema funcional, accesible desde cualquier navegador web sin necesidad "
        "de instalacion local."
    )

    # =========================================================
    # SUBSECCION A: BACKEND EN RENDER.COM
    # =========================================================
    heading(doc, "4.1.5.1. Despliegue del Backend (API REST) en Render.com", level=3)

    body(doc,
        "Render.com es una plataforma de hospedaje en la nube que ofrece una capa gratuita (free tier) para servicios "
        "web que no requieren alta disponibilidad, siendo ideal para proyectos academicos y de demostracion. El backend "
        "del SAA fue contenido en una imagen Docker definida en el archivo Dockerfile ubicado en la raiz del repositorio "
        "del proyecto. Render detecta automaticamente este archivo al importar el repositorio desde GitHub y configura "
        "el ambiente de construccion (build) de manera automatizada."
    )

    body(doc,
        "El proceso de despliegue del backend en Render.com siguio los siguientes pasos:"
    )

    # Lista numerada de pasos render backend
    pasos_render = [
        "Se accedio al portal de Render.com (https://render.com) y se inicio sesion mediante la cuenta de GitHub del desarrollador.",
        "Se selecciono la opcion 'New +' y luego 'Web Service' para crear un nuevo servicio web dinamico capaz de ejecutar la API REST.",
        "Se importo el repositorio de GitHub 'fre14/Sistema_Admision' como fuente del codigo fuente del servicio.",
        "Render detecto automaticamente el archivo Dockerfile en la raiz del repositorio y configuro el lenguaje de construccion como Docker, la rama de despliegue como 'main' y la region como Oregon (US West).",
        "Se asigno el nombre 'Sistema_Admision' al servicio y se selecciono el plan gratuito (Free).",
        "Render inicio el proceso de construccion de la imagen Docker, clonando el repositorio desde GitHub, ejecutando el comando 'dotnet restore' y 'dotnet publish' dentro del contenedor, y generando el artefacto final.",
        "Una vez completada la construccion, Render despliega el contenedor y expone el servicio en el puerto 10000, generando la URL publica: https://sistema-admision-wlii.onrender.com"
    ]

    for i, paso in enumerate(pasos_render, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{paso}")
        set_font(run, size=11)

    doc.add_paragraph()

    # Figura: Crear nuevo servicio en Render
    fn = next_fig()
    insert_image(doc, IMGS["render_crear_servicio"],
        f"Seleccion del tipo de servicio 'Web Service' en el panel de Render.com para el despliegue del backend SAA.",
        fn)

    body(doc,
        "En la pantalla de creacion de nuevo servicio, Render presenta las opciones disponibles en su plataforma. "
        "Para el despliegue del backend del SAA, que es una API REST dinamica contenerizada con Docker, se selecciono "
        "la opcion 'Web Services', la cual permite ejecutar aplicaciones web con logica de servidor, bases de datos en "
        "memoria y endpoints REST, diferenciandose de los 'Static Sites' que solo sirven archivos estaticos."
    )

    # Figura: Seleccion Web Service vs Static Site
    fn = next_fig()
    insert_image(doc, IMGS["render_static_vs_web"],
        f"Pantalla de seleccion entre Static Site y Web Service en Render.com. Se selecciono 'Web Service' para el API REST del SAA.",
        fn)

    body(doc,
        "Render autodetecto la presencia del archivo Dockerfile en el repositorio e identifico que el proyecto "
        "utiliza contenedores Docker para su empaquetado y despliegue. Los campos Name (Sistema_Admision), "
        "Language (Docker), Branch (main) y Region (Oregon US West) fueron configurados para completar el proceso "
        "de configuracion del servicio web."
    )

    # Figura: Configuracion del web service
    fn = next_fig()
    insert_image(doc, IMGS["render_configurar"],
        f"Configuracion del nuevo Web Service en Render.com: nombre del servicio, lenguaje Docker, rama 'main' y region Oregon US West.",
        fn)

    body(doc,
        "Iniciado el proceso de despliegue, Render comenzo a construir la imagen Docker del SAA clonando "
        "el repositorio desde GitHub. Los logs de construccion muestran el commit activo correspondiente al "
        "artefacto de especificaciones SSD (docs(ssd): agregar artefactos OpenSpec - proposal, design, tasks "
        "y source-of-truth del enfoque SSD), evidenciando la trazabilidad entre el control de versiones y el "
        "proceso de despliegue."
    )

    # Figura: Building en curso
    fn = next_fig()
    insert_image(doc, IMGS["render_building"],
        f"Panel de Render.com mostrando el proceso de construccion 'Building' del backend SAA a partir del repositorio de GitHub.",
        fn)

    body(doc,
        "Concluido el proceso de construccion y despliegue, Render emitio el mensaje 'Your service is live' "
        "en el panel de logs del servicio, confirmando que el contenedor Docker del SAA fue ejecutado exitosamente "
        "y que la API REST quedo disponible en la URL publica. Los logs muestran adicionalmente que el servicio "
        "se ejecuta en el puerto 10000, tal como fue configurado en el Dockerfile, y que la aplicacion responde "
        "correctamente a las solicitudes HTTP entrantes."
    )

    # Figura: Servicio live logs
    fn = next_fig()
    insert_image(doc, IMGS["render_live_logs"],
        f"Logs de Render.com confirmando el estado 'Your service is live' del backend SAA. URL disponible: https://sistema-admision-wlii.onrender.com",
        fn)

    body(doc,
        "El backend del SAA quedo accesible de forma permanente en la URL: "
        "https://sistema-admision-wlii.onrender.com. Es importante destacar que, al tratarse de un plan "
        "gratuito, el servicio entra en estado de inactividad (spin down) despues de 15 minutos sin recibir "
        "solicitudes, tardando aproximadamente 50 segundos en reactivarse ante la primera solicitud. Esta "
        "caracteristica es inherente al plan gratuito de Render y no afecta la funcionalidad del sistema para "
        "propositos de evaluacion academica."
    )

    # =========================================================
    # SUBSECCION B: BASE DE DATOS EN SOMEE.COM
    # =========================================================
    heading(doc, "4.1.5.2. Despliegue de la Base de Datos SQL Server en Somee.com", level=3)

    body(doc,
        "Para garantizar la persistencia de los datos del SAA en el entorno de produccion en la nube, se "
        "utilizo la plataforma Somee.com (SmarterASP.NET), que ofrece un plan gratuito de hospedaje de bases "
        "de datos Microsoft SQL Server 2022 Express con 30 MB de capacidad para datos y 30 MB para archivos de "
        "log, suficiente para almacenar los datos del sistema en el contexto de un proyecto academico."
    )

    body(doc,
        "El proceso de creacion de la base de datos en Somee.com siguio los siguientes pasos:"
    )

    pasos_somee = [
        "Se accedio al portal de Somee.com y se creo una cuenta gratuita.",
        "Desde el panel de control, se navego a la seccion 'MS SQL' > 'Databases' > 'Create database'.",
        "Se asigno el nombre 'SAA_AdmisionDB' a la base de datos y se selecciono la zona 'somee.com'.",
        "Somee genero automaticamente un usuario de acceso SQL (fractus_SQLLogin_1) y el servidor de base de datos (SAA_AdmisionDB.mssql.somee.com).",
        "Se copio la cadena de conexion generada por Somee y se configuro como variable de entorno en el backend desplegado en Render.com.",
        "Al iniciarse el backend, el servicio SeedDataService ejecuto automaticamente la creacion de las tablas y la siembra de 500 registros de postulantes de prueba mediante el metodo EnsureCreatedAsync de Entity Framework Core."
    ]

    for i, paso in enumerate(pasos_somee, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{paso}")
        set_font(run, size=11)

    doc.add_paragraph()

    # Figura: Panel Somee con BD creada
    fn = next_fig()
    insert_image(doc, IMGS["somee_bd"],
        f"Panel de administracion de Somee.com mostrando la base de datos SAA_AdmisionDB creada exitosamente con SQL Server 2022 Express.",
        fn)

    body(doc,
        "La figura anterior muestra el panel de administracion de la base de datos SAA_AdmisionDB en Somee.com. "
        "Se puede apreciar la informacion de conexion: SQL Server address (SAA_AdmisionDB.mssql.somee.com), "
        "Login name (fractus_SQLLogin_1), la version del motor (MS SQL 2022 Express) y la capacidad utilizada "
        "(8 MB de datos utilizados de 30 MB disponibles). Esta base de datos actua como el repositorio persistente "
        "de toda la informacion del sistema SAA en el entorno de produccion en la nube."
    )

    body(doc,
        "La cadena de conexion de Somee fue configurada directamente en el archivo appsettings.Production.json "
        "del backend, el cual es leido por Entity Framework Core al iniciarse la aplicacion en el entorno de "
        "produccion. Gracias a esta configuracion, los datos de postulantes, programas academicos, fichas de "
        "postulacion y resultados de examenes son persistentes, es decir, no se pierden cuando el servidor de "
        "Render se reinicia por inactividad."
    )

    # =========================================================
    # SUBSECCION C: FRONTEND EN VERCEL
    # =========================================================
    heading(doc, "4.1.5.3. Despliegue del Frontend (React 19) en Vercel", level=3)

    body(doc,
        "El frontend del SAA, desarrollado con React 19, TypeScript 5.9 y Vite 7, fue desplegado en la plataforma "
        "Vercel, que ofrece un plan gratuito (Hobby) optimizado para aplicaciones frontend modernas basadas en "
        "frameworks JavaScript como React, Next.js y Vite. Vercel detecta automaticamente el framework utilizado "
        "y configura el pipeline de construccion y despliegue sin necesidad de configuracion adicional."
    )

    body(doc,
        "El proceso de despliegue del frontend en Vercel siguio los siguientes pasos:"
    )

    pasos_vercel = [
        "Se accedio al portal de Vercel (https://vercel.com) y se inicio sesion con la cuenta de GitHub del desarrollador.",
        "Se selecciono la opcion 'Add New Project' y se importo el repositorio 'fre14/Sistema_Admision' desde GitHub.",
        "En la pantalla de configuracion del proyecto, se asigno el nombre 'sistema-admision' y se selecciono el equipo 'fnepos' (plan Hobby gratuito).",
        "Se configuro el directorio raiz (Root Directory) apuntando a la carpeta 'SAA.API/frontend', que contiene el codigo fuente del frontend React. Vercel detecto automaticamente el preset 'Vite' para la construccion.",
        "Se hizo clic en 'Deploy' y Vercel ejecuto el pipeline de construccion, instalando las dependencias npm y generando el bundle de produccion optimizado.",
        "Vercel desplego exitosamente el frontend y mostro la pantalla de confirmacion con la vista previa del Portal SAA."
    ]

    for i, paso in enumerate(pasos_vercel, 1):
        p = doc.add_paragraph(style='List Number')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(f"{paso}")
        set_font(run, size=11)

    doc.add_paragraph()

    # Figura: Import desde GitHub en Vercel
    fn = next_fig()
    insert_image(doc, IMGS["vercel_import"],
        f"Importacion del repositorio 'fre14/Sistema_Admision' desde GitHub en Vercel para el despliegue del frontend React del SAA.",
        fn)

    body(doc,
        "La figura anterior muestra la pantalla de creacion de nuevo proyecto en Vercel, donde se puede "
        "observar que el repositorio 'fre14/Sistema_Admision' fue importado exitosamente desde GitHub. El "
        "nombre del proyecto quedo establecido como 'sistema-admision' dentro del equipo 'fnepos', correspondiente "
        "al plan gratuito Hobby de Vercel."
    )

    # Figura: Seleccion Root Directory
    fn = next_fig()
    insert_image(doc, IMGS["vercel_root_dir"],
        f"Seleccion del directorio raiz 'frontend' (SAA.API/frontend) en Vercel, donde Vercel detecto automaticamente el preset de Vite.",
        fn)

    body(doc,
        "La configuracion del directorio raiz es un paso critico en el proceso de despliegue. Dado que el "
        "repositorio contiene multiples proyectos (.NET, C# y React), fue necesario indicarle a Vercel que "
        "el codigo del frontend se encuentra especificamente en la subcarpeta 'frontend' dentro de 'SAA.API'. "
        "Vercel identifico automaticamente el icono de Vite (triangulo amarillo) al seleccionar dicha carpeta, "
        "confirmando que el framework de construccion seria Vite, lo que garantiza una compilacion y optimizacion "
        "correcta del bundle de produccion de React."
    )

    # Figura: Despliegue exitoso
    fn = next_fig()
    insert_image(doc, IMGS["vercel_success"],
        f"Confirmacion de despliegue exitoso en Vercel mostrando la vista previa del Portal SAA con la pantalla de inicio de sesion.",
        fn)

    body(doc,
        "La figura anterior muestra la pantalla de confirmacion de Vercel tras el despliegue exitoso del "
        "frontend del SAA. Se puede apreciar la vista previa de la aplicacion desplegada, que muestra el "
        "Portal SAA con su pantalla de inicio de sesion, con los campos 'Usuario / DNI' y 'Contrasena' "
        "correctamente renderizados. Vercel genero automaticamente una URL publica para acceder al sistema "
        "desde cualquier navegador web sin necesidad de instalacion."
    )

    # =========================================================
    # TABLA RESUMEN INFRAESTRUCTURA
    # =========================================================
    heading(doc, "4.1.5.4. Resumen de la Infraestructura Desplegada en la Nube", level=3)

    body(doc,
        "La siguiente tabla resume la infraestructura completa del Sistema Automatizado de Admision (SAA) "
        "desplegada en servicios de hospedaje gratuitos en la nube, evidenciando el cumplimiento del "
        "requerimiento de despliegue en una capa gratuita de servicios de hosteo."
    )

    # Tabla
    p_tabla_title = doc.add_paragraph()
    p_tabla_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_tabla_title.add_run("Tabla 8")
    set_font(r, bold=True, size=11)

    p_tabla_cap = doc.add_paragraph()
    p_tabla_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p_tabla_cap.add_run("Resumen de la infraestructura del SAA desplegada en servicios gratuitos en la nube")
    set_font(r2, italic=True, size=11)
    p_tabla_cap.paragraph_format.space_after = Pt(6)

    tabla = doc.add_table(rows=1, cols=5)
    tabla.style = 'Table Grid'

    headers = ["Componente", "Plataforma", "Plan", "URL / Host", "Tecnologia"]
    hdr_cells = tabla.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                set_font(run, bold=True, size=10)

    rows_data = [
        ["Backend (API REST)", "Render.com", "Free", "https://sistema-admision-wlii.onrender.com", "ASP.NET Core 10 + Docker"],
        ["Base de Datos", "Somee.com", "Free (30MB)", "SAA_AdmisionDB.mssql.somee.com", "SQL Server 2022 Express"],
        ["Frontend (SPA)", "Vercel", "Hobby (Free)", "URL generada por Vercel", "React 19 + Vite 7"],
        ["Repositorio", "GitHub", "Free", "https://github.com/fre14/Sistema_Admision", "Git + GitHub Actions"],
    ]

    for row_data in rows_data:
        row_cells = tabla.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = val
            for para in row_cells[i].paragraphs:
                for run in para.runs:
                    set_font(run, size=9)

    doc.add_paragraph()

    p_nota_tabla = doc.add_paragraph()
    r_nota = p_nota_tabla.add_run("Nota: Elaboracion propia.")
    set_font(r_nota, size=10, italic=True)
    p_nota_tabla.paragraph_format.space_after = Pt(12)

    body(doc,
        "El despliegue exitoso del SAA en servicios de hospedaje gratuitos en la nube constituye el "
        "resultado final del ciclo de vida del software aplicado en el presente trabajo monografico. "
        "Este resultado demuestra que el sistema desarrollado bajo el enfoque SSD, con arquitectura limpia "
        "y metodologia Scrum, es un producto de software funcional, accesible y desplegable en entornos "
        "reales de produccion, cumpliendo con el Objetivo Especifico OE3 referido a la implementacion "
        "del SAA utilizando Clean Architecture con .NET y React, y cerrando el ciclo de vida con una "
        "evidencia concreta de despliegue en la nube."
    )

    doc.save(OUTPUT)
    print(f"Documento guardado en: {OUTPUT}")
    print(f"Total de figuras insertadas: {fig_counter[0] - 11}")

if __name__ == '__main__':
    build_document()

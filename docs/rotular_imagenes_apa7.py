#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Rotula y corrige en formato APA 7 todas las imagenes del documento
Trabajo_Monografico_SAA_preliminar.docx.
- No elimina ningun parrafo existente
- Solo inserta/corrige los rotulos antes y despues de cada imagen
- Aplica formato APA 7: Figura N en negrita (encima), titulo en italica (debajo), Nota: abajo
"""
import copy
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree

INPUT  = r"C:\Users\usuario\Downloads\Trabajo_Monografico_SAA_preliminar.docx"
OUTPUT = r"C:\Users\usuario\Downloads\Trabajo_Monografico_SAA_FINAL.docx"

# -----------------------------------------------------------------------
# MAPA de imagenes: indice del parrafo con imagen -> datos del rotulo APA7
# Formato: idx_imagen: (numero_figura, titulo_figura, nota)
# -----------------------------------------------------------------------
FIGURA_MAP = {
    # Portada - escudo/logo (no se rotula, se omite)
    4: None,

    # Cap IV - Figura 1: OpenSpec
    348: (1,
          "Estructura de especificaciones automatizadas utilizando la herramienta OpenSpec en el repositorio del SAA.",
          "Elaboracion propia. Captura del explorador de archivos del repositorio fre14/Sistema_Admision mostrando la carpeta openspec/."),

    # Cap IV - Arquitectura Clean Architecture
    409: (2,
          "Diagrama de Arquitectura Limpia (Clean Architecture) aplicada al Sistema Automatizado de Admision (SAA).",
          "Elaboracion propia. Obtenido del repositorio del proyecto SAA, carpeta docs/."),

    # Cap IV - Estructura del repositorio (dos imagenes juntas)
    472: (3,
          "Estructura del repositorio de GitHub fre14/Sistema_Admision evidenciando el control de versiones del SAA.",
          "Elaboracion propia. Captura del repositorio en GitHub.com."),

    473: (4,
          "Vista detallada de la estructura de carpetas del proyecto SAA en el repositorio de control de versiones.",
          "Elaboracion propia. Captura del repositorio en GitHub.com."),

    # Cap IV - Interfaces del sistema
    488: (5,
          "Pantalla de inicio de sesion del Portal SAA desplegado en produccion.",
          "Elaboracion propia. Captura de la interfaz del Sistema Automatizado de Admision (SAA) en produccion."),

    491: (6,
          "Dashboard del Administrador en el Portal SAA: gestion de candidatos, procesamiento de resultados y reporte general.",
          "Elaboracion propia. Captura de la interfaz del modulo de administracion del SAA."),

    494: (7,
          "Dashboard del Postulante en el Portal SAA: visualizacion personalizada del estado y resultado de admision.",
          "Elaboracion propia. Captura de la interfaz del modulo de postulante del SAA."),

    # Cap IV - Pruebas unitarias
    500: (8,
          "Resultado de la ejecucion de las pruebas unitarias del SAA en el Explorador de Pruebas de Visual Studio 2022.",
          "Elaboracion propia. Captura del Explorador de Pruebas de Visual Studio 2022 mostrando las pruebas unitarias aprobadas del SAA."),

    # Cap IV - Pruebas de integracion
    559: (9,
          "Resultado de la ejecucion de las pruebas de integracion del SAA utilizando xUnit y base de datos InMemory.",
          "Elaboracion propia. Captura del Explorador de Pruebas de Visual Studio 2022 mostrando las pruebas de integracion aprobadas del SAA."),

    # Cap IV - Reporte de cobertura
    567: (10,
          "Reporte de cobertura de codigo del SAA generado por Coverlet, mostrando porcentaje de lineas, ramas e instrucciones cubiertas.",
          "Elaboracion propia. Reporte HTML generado por la herramienta Coverlet en el proyecto SAA.Tests."),

    # Cap IV Despliegue - ya rotuladas, pero las ajustamos
    582: (11,
          "Seleccion del tipo de servicio 'Web Service' en el panel de Render.com para el despliegue del backend SAA.",
          "Elaboracion propia. Captura del portal Render.com durante el proceso de despliegue del backend del SAA."),

    587: (12,
          "Pantalla de seleccion entre Static Site y Web Service en Render.com. Se selecciono 'Web Service' para la API REST del SAA.",
          "Elaboracion propia. Captura del portal Render.com, seccion de creacion de nuevo servicio web."),

    592: (13,
          "Configuracion del nuevo Web Service en Render.com: nombre del servicio, lenguaje Docker, rama main y region Oregon US West.",
          "Elaboracion propia. Captura del formulario de configuracion del servicio web del SAA en Render.com."),

    597: (14,
          "Panel de Render.com mostrando el proceso de construccion 'Building' del backend SAA a partir del repositorio de GitHub.",
          "Elaboracion propia. Captura del panel de despliegue de Render.com durante la construccion de la imagen Docker del SAA."),

    602: (15,
          "Logs de Render.com confirmando el estado 'Your service is live' del backend SAA disponible en https://sistema-admision-wlii.onrender.com.",
          "Elaboracion propia. Captura de los logs de aplicacion del servicio web del SAA en Render.com."),

    616: (16,
          "Panel de administracion de Somee.com mostrando la base de datos SAA_AdmisionDB creada exitosamente con SQL Server 2022 Express.",
          "Elaboracion propia. Captura del panel de control de Somee.com con los detalles de conexion de la base de datos del SAA."),

    632: (17,
          "Importacion del repositorio fre14/Sistema_Admision desde GitHub en Vercel para el despliegue del frontend React del SAA.",
          "Elaboracion propia. Captura del portal Vercel durante la configuracion del nuevo proyecto del frontend SAA."),

    637: (18,
          "Seleccion del directorio raiz 'frontend' (SAA.API/frontend) en Vercel donde Vercel detecto automaticamente el preset de Vite.",
          "Elaboracion propia. Captura de la seleccion del directorio raiz del frontend SAA en Vercel."),

    642: (19,
          "Confirmacion de despliegue exitoso en Vercel mostrando la vista previa del Portal SAA con la pantalla de inicio de sesion.",
          "Elaboracion propia. Captura de la pantalla de confirmacion de despliegue del frontend SAA en Vercel."),

    # ANEXOS
    784: (20,
          "Diagrama de Arquitectura Limpia (Clean Architecture) del SAA — Anexo 4.",
          "Elaboracion propia."),

    787: (21,
          "Diagrama Entidad-Relacion (DER) del Sistema Automatizado de Admision (SAA) — Anexo 4.",
          "Elaboracion propia."),

    793: (22,
          "Prototipo 1: Pantalla de Login del Portal SAA — Anexo 4.",
          "Elaboracion propia."),

    795: (23,
          "Prototipo 2: Dashboard del Administrador del Portal SAA — Anexo 4.",
          "Elaboracion propia."),

    797: (24,
          "Prototipo 3: Dashboard del Postulante del Portal SAA — Anexo 4.",
          "Elaboracion propia."),
}

def make_run(para, text, bold=False, italic=False, size=11, color=None):
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def make_para_after(doc, ref_para, text, align=WD_ALIGN_PARAGRAPH.CENTER,
                    bold=False, italic=False, size=11, space_before=0, space_after=4):
    """Inserta un nuevo parrafo justo despues de ref_para en el documento."""
    new_para = OxmlElement('w:p')
    ref_para._element.addnext(new_para)
    
    # Importar el parrafo correctamente usando la API de docx
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_para, ref_para._element.getparent())
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def make_para_before(doc, ref_para, text, align=WD_ALIGN_PARAGRAPH.CENTER,
                     bold=False, italic=False, size=11, space_before=12, space_after=2):
    """Inserta un nuevo parrafo justo antes de ref_para en el documento."""
    new_para = OxmlElement('w:p')
    ref_para._element.addprevious(new_para)
    
    from docx.text.paragraph import Paragraph
    p = Paragraph(new_para, ref_para._element.getparent())
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p

def process_document():
    doc = Document(INPUT)
    paras = doc.paragraphs

    # Construir conjunto de indices que YA tienen rotulo tipo "Figura N"
    # para no duplicar
    existing_figura_labels = set()
    for i, p in enumerate(paras):
        txt = p.text.strip()
        if txt.startswith('Figura ') and len(txt) < 15:
            try:
                num = int(txt.replace('Figura', '').strip())
                existing_figura_labels.add(num)
            except:
                pass

    print(f"Figuras ya etiquetadas en el doc: {sorted(existing_figura_labels)}")

    # Procesar cada imagen en el FIGURA_MAP
    # IMPORTANTE: procesar en orden INVERSO para que los indices no se desplacen
    # al insertar parrafos nuevos
    indices_ordenados = sorted(FIGURA_MAP.keys(), reverse=True)

    for idx in indices_ordenados:
        data = FIGURA_MAP[idx]
        if data is None:
            print(f"  [SKIP] parrafo {idx}: imagen de portada, se omite rotulo")
            continue

        num_fig, titulo, nota = data
        img_para = paras[idx]

        print(f"  [PROC] parrafo {idx}: Figura {num_fig} -> {titulo[:60]}...")

        # ---------------------------------------------------------
        # 1. INSERTAR NOTA debajo de la imagen (se inserta primero
        #    porque trabajamos en orden inverso: primero el de abajo)
        # ---------------------------------------------------------
        # Verificar si el parrafo siguiente ya tiene "Nota:"
        next_text = paras[idx+1].text.strip() if idx+1 < len(paras) else ''
        already_has_nota = next_text.startswith('Nota:') or next_text.startswith('Nota.')

        if not already_has_nota:
            make_para_after(doc, img_para,
                            f"Nota: {nota}",
                            align=WD_ALIGN_PARAGRAPH.CENTER,
                            italic=True, size=10,
                            space_before=2, space_after=10)

        # ---------------------------------------------------------
        # 2. INSERTAR TITULO en italica debajo de imagen (encima de nota)
        # ---------------------------------------------------------
        # Revisar si el parrafo siguiente (despues de insertar nota) ya es el titulo
        next_text2 = paras[idx+1].text.strip() if idx+1 < len(paras) else ''
        # Si ya existe un titulo (no empieza con Figura, no es Nota y no esta vacio)
        already_has_titulo = (
            next_text2 and
            not next_text2.startswith('Figura') and
            not next_text2.startswith('Nota:') and
            len(next_text2) > 10
        )

        if not already_has_titulo:
            make_para_after(doc, img_para,
                            titulo,
                            align=WD_ALIGN_PARAGRAPH.CENTER,
                            italic=True, size=11,
                            space_before=2, space_after=2)

        # ---------------------------------------------------------
        # 3. INSERTAR "Figura N" en negrita ANTES de la imagen
        # ---------------------------------------------------------
        # Verificar si ya existe el label "Figura N" antes
        prev_text = paras[idx-1].text.strip() if idx > 0 else ''
        already_has_label = prev_text.startswith(f'Figura {num_fig}')

        if not already_has_label:
            make_para_before(doc, img_para,
                             f"Figura {num_fig}",
                             align=WD_ALIGN_PARAGRAPH.CENTER,
                             bold=True, size=11,
                             space_before=12, space_after=2)

    doc.save(OUTPUT)
    print(f"\nDocumento guardado en: {OUTPUT}")
    print(f"Total de figuras procesadas: {len([k for k,v in FIGURA_MAP.items() if v is not None])}")

if __name__ == '__main__':
    process_document()
